import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

base_dir = r"c:\Users\hp\Downloads\ahuja images\Universal_Solar_Stories_Improved_Package\universal_solar_stories_website_improved"
assets_dir = os.path.join(base_dir, "assets", "images")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_english_doc():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("☀️ Universal Electronics Ethiopia\nSolar Stories Package")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(11, 47, 143)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Official Content & TikTok Short-Form Video Production Suite (9:16 Vertical Format)\nTefera Business Center, Habte Giyorgis, Addis Ababa, Ethiopia | universalelectronicset.com")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(90, 100, 117)

    doc.add_paragraph() # spacing

    stories = [
        {
            "num": "01",
            "title": "The 19-Year-Old Who Found Electricity in Light",
            "category": "History",
            "factline": "1839 • An Electrochemical Experiment • The Photovoltaic Effect",
            "img": "01-becquerel-tiktok.jpg",
            "hook": "“What were you doing at 19? Edmond Becquerel was discovering that light could create electricity.”",
            "article": [
                "Imagine a laboratory in Paris in 1839. There are no solar farms, no rooftop panels and no lithium batteries. Electricity itself is still a young field of science. In that environment, 19-year-old Alexandre-Edmond Becquerel was working in the laboratory of his father, physicist Antoine César Becquerel.",
                "Edmond placed metal electrodes in a conductive liquid and exposed part of the experiment to light. He observed that the electrical response changed when light reached the system. This was an early demonstration of what became known as the photovoltaic effect: light can cause a material or device to produce voltage and current.",
                "The experiment did not immediately create a practical solar panel. The current was small, the materials were different from modern silicon cells, and scientists still needed decades of work to understand semiconductors. But the idea was revolutionary: sunlight was not only warmth and illumination—it could be converted directly into electricity."
            ],
            "script": [
                ("0–04 sec", "Open on a dark, old laboratory. Text: 'Paris, 1839.'", "What were you doing at 19?"),
                ("04–12 sec", "A young scientist prepares two electrodes in liquid.", "Edmond Becquerel was experimenting with light and electricity in his father’s laboratory."),
                ("12–22 sec", "A beam of sunlight reaches one side of the experiment; a meter needle moves.", "When light touched the setup, the electrical response changed."),
                ("22–32 sec", "Close-up animation: photons strike the device and current begins to flow.", "He had observed an early form of the photovoltaic effect—light creating electrical current."),
                ("32–46 sec", "Transition from the 1839 experiment to early solar cells.", "It was not yet a modern solar panel. The power was tiny, and the science needed many more discoveries."),
                ("46–61 sec", "Modern rooftop panels, solar farms and a satellite appear.", "But that small laboratory observation became the starting point for technology used on Earth and in space.")
            ],
            "caption": "At only 19, Edmond Becquerel observed that light could produce an electrical current. His 1839 experiment did not look like today’s solar panel, but it revealed the principle that helped make modern solar power possible.",
            "hashtags": "#UniversalElectronics #SolarHistory #SolarEnergy #ScienceStory #Innovation #RenewableEnergy #Ethiopia"
        },
        {
            "num": "02",
            "title": "There Are No Electric Poles in Space",
            "category": "Space Science",
            "factline": "Solar Arrays • Rechargeable Lithium Batteries • Orbit Eclipse Power",
            "img": "02-space-tiktok.jpg",
            "hook": "“There are no electric poles in space—so how does a satellite stay powered?”",
            "article": [
                "A satellite may look silent from Earth, but inside it is constantly working. It may be collecting weather data, transmitting television signals, supporting navigation, photographing the planet or sending scientific measurements back to a control centre. Every one of those tasks needs electricity.",
                "For many spacecraft operating near Earth or closer to the Sun, the practical answer is solar power. Solar cells mounted on panels convert sunlight into electricity. The electricity runs the spacecraft’s instruments and also charges onboard batteries.",
                "The batteries are essential because a spacecraft does not always have direct sunlight. A satellite orbiting Earth can pass into Earth’s shadow. During these periods, stored energy keeps computers, heaters, communications and control systems operating."
            ],
            "script": [
                ("0–04 sec", "Satellite floats above Earth. Empty space surrounds it.", "There are no electric poles in space."),
                ("04–10 sec", "Zoom toward the satellite’s dark solar arrays.", "So how does a satellite keep its cameras, computers and communication systems running?"),
                ("10–21 sec", "Panels rotate into sunlight and begin glowing subtly.", "Solar cells convert sunlight into electricity while the spacecraft is exposed to the Sun."),
                ("21–33 sec", "Animated energy flows from panels to onboard systems and battery.", "Part of that electricity powers the mission. Another part charges rechargeable batteries."),
                ("33–45 sec", "Satellite enters Earth’s shadow; solar glow disappears.", "When the spacecraft moves into darkness, the batteries take over.")
            ],
            "caption": "No electric poles. No fuel station. No extension cable. Many spacecraft use solar arrays to generate electricity and batteries to stay powered when they pass into shadow.",
            "hashtags": "#UniversalElectronics #SpaceTechnology #SolarPower #Satellite #BatteryStorage #Ethiopia"
        },
        {
            "num": "03",
            "title": "Solar Energy Existed Before Solar Panels",
            "category": "Human Ingenuity",
            "factline": "Traditional Solar Thermal • Crop Drying • Natural Architecture",
            "img": "03-before-panels-tiktok.jpg",
            "hook": "“Solar energy is not a modern invention. Solar panels are only the modern tool.”",
            "article": [
                "When people hear 'solar energy,' they usually imagine a dark panel on a roof. But humans were using the Sun’s energy thousands of years before electricity.",
                "Buildings were placed to use daylight and seasonal warmth. Communities dried clothes, crops, fruit, meat and other foods in the Sun. In Ethiopia, coffee, grains and spices are often dried using solar heat and moving air. A raised drying bed is a system designed to use energy from the Sun.",
                "Modern photovoltaic systems add a powerful new step: converting light into electricity to run lights, computers and appliances."
            ],
            "script": [
                ("0–05 sec", "Modern rooftop solar panel appears.", "You may think solar energy began with this."),
                ("05–13 sec", "Panel fades into an ancient sunlit settlement.", "But people used solar energy long before anyone created a solar cell."),
                ("39–53 sec", "Ethiopian coffee cherries dry on raised beds.", "In Ethiopia, the idea is still easy to recognise in coffee, grain and agricultural drying.")
            ],
            "caption": "Solar energy did not begin with rooftop panels. People have long used sunlight for natural lighting, warmth, drying and preservation.",
            "hashtags": "#UniversalElectronics #SolarStory #EthiopianCoffee #TraditionalKnowledge #Ethiopia"
        },
        {
            "num": "04",
            "title": "How Powerful Is the Sun?",
            "category": "Solar Scale",
            "factline": "173,000 Terawatts • Water Cycle • Clean Solar Electricity",
            "img": "04-sun-power-tiktok.jpg",
            "hook": "“The Sun sends Earth more energy than humanity knows how to use—but most of it passes us by.”",
            "article": [
                "The Sun’s energy influences Earth’s temperature, weather, water cycle, winds and ocean circulation. NASA EarthData explains that Earth receives more than 10,000 times as much energy from the Sun as the entire planet uses.",
                "A solar panel does not need to capture all sunlight. It only needs to convert a useful portion of the light falling on its surface into clean electricity."
            ],
            "script": [
                ("0–05 sec", "A tiny Earth appears beside the Sun.", "How powerful is the Sun?"),
                ("25–39 sec", "Large number visual: 10,000×.", "NASA says Earth receives more than ten thousand times the energy humanity uses.")
            ],
            "caption": "Earth receives more than 10,000 times the energy humanity uses, according to NASA EarthData.",
            "hashtags": "#UniversalElectronics #SolarFacts #TheSun #RenewableEnergy #Ethiopia"
        },
        {
            "num": "05",
            "title": "Why Are Solar Panels Blue or Black?",
            "category": "Everyday Science",
            "factline": "Polycrystalline Blue • Monocrystalline Black • Anti-Reflective Coating",
            "img": "05-panel-colour-tiktok.jpg",
            "hook": "“Why are most solar panels dark blue or black instead of white, red or green?”",
            "article": [
                "A solar panel must allow as much light as possible to enter its cells. Silicon naturally reflects light when untreated. Manufacturers use textured surfaces and anti-reflection coatings to reduce reflection, resulting in a dark blue or black appearance.",
                "Multicrystalline modules look blue with a patterned surface. Monocrystalline modules appear darker and uniform. Colour alone does not prove quality—cell design, manufacturing standards, and installation quality matter far more."
            ],
            "script": [
                ("0–05 sec", "Blue and black panels appear side by side.", "Which one is better: blue or black?"),
                ("40–53 sec", "A large red X appears over 'colour = quality.'", "But colour alone does not tell you which panel is better.")
            ],
            "caption": "Dark blue and black panels are designed to absorb light and reduce reflection, but colour alone is not a quality certificate.",
            "hashtags": "#UniversalElectronics #SolarPanels #SolarEducation #TechExplained #Ethiopia"
        },
        {
            "num": "06",
            "title": "What If the Sun Suddenly Disappeared?",
            "category": "Thought Experiment",
            "factline": "About 8 Minutes • Darkness • Earth Leaves Orbit",
            "img": "06-sun-disappears-tiktok.jpg",
            "hook": "“If the Sun disappeared right now, you would not notice immediately.”",
            "article": [
                "This cannot physically happen, but it is a powerful thought experiment. Sunlight takes roughly eight minutes to travel from the Sun to Earth. Changes in gravity also propagate at the speed of light.",
                "If the Sun vanished instantly, Earth would continue seeing sunlight and following its orbit for about eight minutes. After that delay, the sky would go dark and Earth would continue moving forward along a straight tangent path."
            ],
            "script": [
                ("0–04 sec", "Normal morning. The Sun vanishes silently.", "If the Sun disappeared right now, you would not notice immediately."),
                ("04–14 sec", "Clock begins counting toward eight minutes.", "The sunlight already travelling toward Earth would continue arriving for about eight minutes."),
                ("24–36 sec", "At eight minutes the world becomes dark.", "Then the last sunlight would pass, and the sky would go dark.")
            ],
            "caption": "If the Sun magically disappeared, Earth would still receive its existing light and gravitational influence for about eight minutes.",
            "hashtags": "#UniversalElectronics #WhatIf #TheSun #SpaceScience #Ethiopia"
        },
        {
            "num": "07",
            "title": "Could You Live Completely Without Grid Electricity?",
            "category": "Audience Challenge",
            "factline": "Know Your Loads • Store Energy • Plan For Low Generation",
            "img": "07-off-grid-tiktok.jpg",
            "hook": "“Would you disconnect from the grid today and trust only solar power?”",
            "article": [
                "Living without grid electricity sounds simple: put panels on the roof, add a battery and stop worrying about power cuts. In reality, a dependable off-grid system begins with load analysis.",
                "Separating critical loads (lights, Wi-Fi, refrigeration) from high-power loads (electric cookers, water heaters) changes the required system size."
            ],
            "script": [
                ("0–05 sec", "A hand switches off the main grid breaker.", "Would you disconnect from the grid today and depend completely on solar?"),
                ("05–15 sec", "House remains lit; panel, inverter and battery appear.", "It is possible—but it is not as simple as buying panels and a battery.")
            ],
            "caption": "Could you live completely without grid electricity? A reliable off-grid system starts with load analysis.",
            "hashtags": "#UniversalElectronics #OffGridLiving #SolarChallenge #EnergyIndependence #Ethiopia"
        }
    ]

    for item in stories:
        h2 = doc.add_paragraph()
        run_h2 = h2.add_run(f"Story {item['num']}: {item['title']}")
        run_h2.font.name = 'Calibri'
        run_h2.font.size = Pt(16)
        run_h2.font.bold = True
        run_h2.font.color.rgb = RGBColor(11, 47, 143)

        p_meta = doc.add_paragraph()
        run_meta = p_meta.add_run(f"Category: {item['category']} | {item['factline']}")
        run_meta.font.name = 'Calibri'
        run_meta.font.size = Pt(9)
        run_meta.font.bold = True
        run_meta.font.color.rgb = RGBColor(22, 73, 198)

        # Image Insertion
        img_path = os.path.join(assets_dir, item['img'])
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(12)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(2.5)) # 9:16 portrait scale

        # Hook Box Table
        table_hook = doc.add_table(rows=1, cols=1)
        table_hook.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_hook = table_hook.cell(0, 0)
        set_cell_background(cell_hook, "FFF9ED")
        p_hook = cell_hook.paragraphs[0]
        p_hook.paragraph_format.space_before = Pt(6)
        p_hook.paragraph_format.space_after = Pt(6)
        run_hk_lbl = p_hook.add_run("TIKTOK OPENING HOOK:\n")
        run_hk_lbl.font.size = Pt(8.5)
        run_hk_lbl.font.bold = True
        run_hk_lbl.font.color.rgb = RGBColor(179, 123, 0)
        run_hk_txt = p_hook.add_run(item['hook'])
        run_hk_txt.font.size = Pt(11)
        run_hk_txt.font.bold = True
        run_hk_txt.font.color.rgb = RGBColor(6, 24, 69)

        doc.add_paragraph() # space

        # Article Paragraphs
        for p_txt in item['article']:
            p_art = doc.add_paragraph()
            r_art = p_art.add_run(p_txt)
            r_art.font.name = 'Calibri'
            r_art.font.size = Pt(10.5)
            p_art.paragraph_format.space_after = Pt(6)

        # Script Table
        p_sc_head = doc.add_paragraph()
        r_sc_head = p_sc_head.add_run("📱 90-Second TikTok Video Script & Teleprompter")
        r_sc_head.font.bold = True
        r_sc_head.font.size = Pt(11)
        r_sc_head.font.color.rgb = RGBColor(11, 47, 143)

        table_sc = doc.add_table(rows=1, cols=3)
        table_sc.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table_sc.rows[0].cells
        hdr_cells[0].text = "Timing"
        hdr_cells[1].text = "Visual Direction"
        hdr_cells[2].text = "Narration"

        for cell in hdr_cells:
            set_cell_background(cell, "EEF3FF")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(6, 24, 69)

        for step in item['script']:
            row_cells = table_sc.add_row().cells
            row_cells[0].text = step[0]
            row_cells[1].text = step[1]
            row_cells[2].text = f'"{step[2]}"'
            for idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        if idx == 0:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(11, 47, 143)
                        elif idx == 2:
                            r.font.bold = True

        doc.add_paragraph()

        # Caption
        p_cap = doc.add_paragraph()
        r_cap_lbl = p_cap.add_run("Social Media Caption:\n")
        r_cap_lbl.font.bold = True
        r_cap_lbl.font.size = Pt(9.5)
        r_cap_txt = p_cap.add_run(item['caption'] + "\n" + item['hashtags'])
        r_cap_txt.font.size = Pt(9.5)
        r_cap_txt.font.italic = True

        doc.add_paragraph("--------------------------------------------------------------------------------")

    output_path = os.path.join(base_dir, "Universal_Solar_Stories_English.docx")
    doc.save(output_path)
    print("Created:", output_path)

def create_amharic_doc():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("☀️ ዩኒቨርሳል ኤሌክትሮኒክስ ኢትዮጵያ\nየሶላር ታሪኮች ስብስብ")
    run_title.font.name = 'Nyala'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(11, 47, 143)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("የቲክቶክ፣ ሪልስ እና ሾርትስ የቪዲዮ ይዘቶችና ስክሪፕቶች (9:16 ቨርቲካል ፎርማት)\nተፈራ ቢዝነስ ሴንተር፣ ሀብተ ጊዮርጊስ፣ አዲስ አበባ፣ ኢትዮጵያ | universalelectronicset.com")
    run_sub.font.name = 'Nyala'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(90, 100, 117)

    doc.add_paragraph() # spacing

    am_stories = [
        {
            "num": "01",
            "title": "በብርሃን ውስጥ ኤሌክትሪክን ያገኘው የ19 ዓመቱ ወጣት",
            "category": "ታሪክ",
            "factline": "1839 • የላቦራቶሪ ሙከራ • የፎቶቮልታይክ ግኝት (Photovoltaic Effect)",
            "img": "01-becquerel-tiktok.jpg",
            "hook": "“በ19 ዓመታችሁ ምን ታደርጉ ነበር? ኤድሞንድ ቤክሬል ብርሃን ኤሌክትሪክ ማመንጨት እንደሚችል ሲያገኝ ነበር!”",
            "article": [
                "በ1839 በፓሪስ የሚገኝ አንድ የድሮ የሳይንስ ላቦራቶሪ ያስቡ። በዚያ ዘመን ምንም ዓይነት የሶላር ፓነሎች፣ የሊቲየም ባትሪዎች ወይም ግዙፍ የኤሌክትሪክ ጣቢያዎች አልነበሩም። ኤሌክትሪክ ራሱ ገና አዲስና አጃኢብ የሚባል የሳይንስ መስክ ነበር። በዚያ አስደናቂ ሁኔታ ውስጥ፣ የ19 ዓመቱ አሌክሳንደር ኤድሞንድ ቤክሬል የታዋቂው ፊዚሲስት አባቱ ላቦራቶሪ ውስጥ ምርምር ያደርግ ነበር።",
                "ኤድሞንድ የብረት ኤሌክትሮዶችን በፈሳሽ ውስጥ አድርጎ የፀሐይ ብርሃን በሙከራው ላይ ሲያርፍበት፣ የኤሌክትሪክ ፍሰቱ በከፍተኛ ሁኔታ እንደተቀየረና ኃይል እንደተፈጠረ አስተዋለ። ይህ ታሪካዊ ግኝት Photovoltaic Effect (የፎቶቮልታይክ ክስተት) በመባል ታወቀ—ይህም ብርሃን ቁሳቁሶችን በመንካት ኤሌክትሪክ እንዲያመነጩ የማድረግ ተፈጥሯዊ ህግ ነው።",
                "የቤክሬል ሙከራ ያኔ ወዲያውኑ የቤት ላይ ሶላር ፓነል አልፈጠረም። የሚመነጨው ኤሌክትሪክ በጣም ትንሽ ነበር። ነገር ግን አስተሳሰቡ አብዮታዊ ነበር፡ የፀሐይ ብርሃን የሙቀትና የብርሃን ምንጭ ብቻ ሳይሆን በቀጥታ ወደ ኤሌክትሪክ ኃይል የሚቀየር ታላቅ የተፈጥሮ ጸጋ መሆኑ ተረጋገጠ።"
            ],
            "script": [
                ("0–04 ሰከንድ", "የድሮ ላቦራቶሪ። ጽሑፍ: 'ፓሪስ፣ 1839'", "በ19 ዓመታችሁ ምን ታደርጉ ነበር?"),
                ("04–12 ሰከንድ", "ወጣቱ ሳይንቲስት በፈሳሽ ውስጥ ኤሌክትሮዶችን ያዘጋጃል።", "ኤድሞንድ ቤክሬል በአባቱ ላቦራቶሪ ውስጥ በብርሃንና በኤሌክትሪክ ይሞክር ነበር።"),
                ("12–22 ሰከንድ", "የፀሐይ ጨረር በሙከራው ላይ ሲያርፍ የቆጣሪው መርፌ ይንቀሳቀሳል።", "ብርሃኑ ሙከራውን ሲነካው የኤሌክትሪክ እንቅስቃሴው ተቀየረ።"),
                ("22–32 ሰከንድ", "ፎቶኖች ሲያርፉ ኤሌክትሪክ ሲፈስ የሚያሳይ አኒሜሽን።", "ብርሃን ኤሌክትሪክ ማመንጨት እንደሚችል የመጀመሪያውን ማረጋገጫ አየ።"),
                ("32–46 ሰከንድ", "ከ1839 ሙከራ ወደ ዘመናዊ የሶላር ሴሎች ሽግግር።", "ያኔ ወዲያውኑ የሶላር ፓነል አልተሰራም፣ ነገር ግን ትልቅ በር ከፈተ።"),
                ("46–61 ሰከንድ", "የቤት ላይ ሶላር ፓነሎች እና ሳተላይቶች።", "ያች ትንሽ ላቦራቶሪ ዛሬ በምድርና በህዋ ላይ ለምንጠቀመው ቴክኖሎጂ መነሻ ሆነች።")
            ],
            "caption": "በ19 ዓመቱ ኤድሞንድ ቤክሬል ብርሃን ኤሌክትሪክ ማመንጨት እንደሚችል አገኘ! የ1839 ሙከራው ለዛሬው የሶላር ፓነል መሰረት ሆነ። ታላላቅ ቴክኖሎጂዎች ብዙውን ጊዜ የሚጀምሩት ከአንድ ትንሽ ምልከታ ነው።",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #የሶላርታሪክ #ኢትዮጵያ #ሳይንስ #የሶላርኃይል #አዲስአበባ"
        },
        {
            "num": "02",
            "title": "በህዋ (Space) ውስጥ የኤሌክትሪክ ፖሎች የሉም",
            "category": "ህዋ",
            "factline": "የሶላር ፓነሎች • ሊቲየም ባትሪዎች • በጨለማ ወቅት ኃይል መስጠት",
            "img": "02-space-tiktok.jpg",
            "hook": "“በህዋ ውስጥ የኤሌክትሪክ ፖሎች የሉም—ታዲያ ሳተላይቶች እንዴት ኃይል ያገኛሉ?”",
            "article": [
                "ሳተላይት ከምድር ሲታይ በህዋ ውስጥ ጸጥ ብሎ የሚንሳፈፍ ይመስላል፤ ነገር ግን በውስጡ ሳያቋርጥ ይሰራል! የአየር ሁኔታ መረጃዎችን ይሰበስባል፣ የቴሌቪዥንና ስልክ ምልክቶችን ያስተላልፋል፣ የጂፒኤስ መሪዎችን ይመራል፣ እንዲሁም ሳይንሳዊ ምርምሮችን ያከናውናል። እያንዳንዱ ተግባር አስተማማኝ ኤሌክትሪክ ይፈልጋል።",
                "በምድር ዙሪያ ለሚሰሩ አብዛኞቹ ሳተላይቶች መፍትሄው የፀሐይ ኃይል ነው። በሳተላይቱ ክንፎች ላይ የተገጠሙት የሶላር ሴሎች የፀሐይ ብርሃንን ወደ ኤሌክትሪክ በመቀየር የሳተላይቱን መሳሪያዎች ያሰራሉ፤ እንዲሁም ውስጣዊ የሊቲየም ባትሪዎችን ይሞላሉ።",
                "ባትሪዎቹ በጣም ወሳኝ ናቸው ምክንያቱም ሳተላይቱ በምድር ዙሪያ ሲሽከረከር በምድር ጥላ ስር (Night side) ያልፋል። በዚያ ጨለማ ወቅት የሶላር ፓነሎቹ ብርሃን አያገኙም፤ ታዲያ በባትሪ የተቀመጠው ኃይል ሳተላይቱ ሳይጠፋ ስራውን እንዲቀጥል ያደርጋል!"
            ],
            "script": [
                ("0–04 ሰከንድ", "ሳተላይት በህዋ ላይ ይንሳፈፋል።", "በህዋ ውስጥ የኤሌክትሪክ ፖሎች የሉም።"),
                ("04–10 ሰከንድ", "ወደ ሶላር ፓነሎች ማቀረብ።", "ታዲያ ሳተላይቶች ካሜራዎቻቸውንና ኮምፒውተሮቻቸውን እንዴት ያስኬዳሉ?"),
                ("10–21 ሰከንድ", "ፓነሎች ወደ ፀሐይ ዞረው ሲበሩ።", "የሶላር ሴሎች የፀሐይ ብርሃንን ወደ ኤሌክትሪክ ይቀይራሉ።"),
                ("21–33 ሰከንድ", "ኃይል ወደ ባትሪ ሲፈስ።", "ከፊሉ ኤሌክትሪክ ሳተላይቱን ሲያሰራ፣ ከፊሉ ባትሪ ይሞላል።"),
                ("33–45 ሰከንድ", "ሳተላይት ወደ ምድር ጥላ ሲገባ።", "ሳተላይቱ ወደ ጨለማ ሲገባ በባትሪ የተቀመጠው ኃይል ስራውን ይቀጥላል።")
            ],
            "caption": "በህዋ ውስጥ የኤሌክትሪክ ፖል የለም! ሳተላይቶች በሶላር ፓነሎችና ባትሪዎች ኃይል በማመንጨትና በማስቀመጥ ለዓመታት ይሰራሉ።",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #የሶላርታሪክ #ህዋ #ቴክኖሎጂ #ኢትዮጵያ"
        },
        {
            "num": "03",
            "title": "የፀሐይ ኃይል ከሶላር ፓነሎች በፊትም ነበረ",
            "category": "የሰው ልጅ ፈጠራ",
            "factline": "የባህላዊ የፀሐይ አጠቃቀም • ቡና ማድረቅ • ተፈጥሯዊ አየር",
            "img": "03-before-panels-tiktok.jpg",
            "hook": "“የፀሐይ ኃይል አዲስ ፈጠራ አይደለም። የሶላር ፓነሎች ዘመናዊ መሣሪያዎች ናቸው!”",
            "article": [
                "ሰዎች 'የሶላር ኃይል' ሲባል ብዙውን ጊዜ ጣሪያ ላይ የሚቀመጠውን ጥቁር የሶላር ፓነል ያስባሉ። ነገር ግን የሰው ልጅ ኤሌክትሪክ ከመፈጠሩ በሺዎች ከሚቆጠሩ ዓመታት በፊት የፀሐይን ኃይል ይጠቀም ነበር!",
                "ጥንታውያን ህንጻዎች የፀሐይን ብርሃንና ሙቀት በጥበብ እንዲቀበሉ ተደርገው ይሰሩ ነበር። አባቶች እህል፣ ፍራፍሬ፣ ስጋና ልብሶችን በፀሐይ ያደርቁ ነበር። በሃገራችን ኢትዮጵያም ቡና፣ እህልና ቅመማ ቅመሞች በፀሐይ ብርሃንና በአየር እንቅስቃሴ ይደርቃሉ።"
            ],
            "script": [
                ("0–05 ሰከንድ", "ዘመናዊ የሶላር ፓነል ይታያል።", "የሶላር ኃይል የጀመረው በዚህ ይመስላችሁ ይሆናል።"),
                ("05–13 ሰከንድ", "ወደ ጥንታዊ መንደር ይሸጋገራል።", "ነገር ግን የሰው ልጅ የሶላር ሴል ከመፈጠሩ በፊትም የፀሐይ ኃይልን ይጠቀም ነበር።"),
                ("39–53 ሰከንድ", "የኢትዮጵያ ቡና ሲደርቅ።", "በኢትዮጵያም በቡናና በእህል ማድረቅ ላይ ይህን ባህል እናየዋለን።")
            ],
            "caption": "የፀሐይ ኃይል በሶላር ፓነል አልጀመረም። ሰዎች ከጥንት ጀምሮ ለብርሃን፣ ለማሞቅና ቡና ለማድረቅ ይጠቀሙበት ነበር።",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #የሶላርታሪክ #የኢትዮጵያቡና #ተፈጥሮ #ኢትዮጵያ"
        },
        {
            "num": "04",
            "title": "ፀሐይ ምን ያህል ኃይለኛ ናት?",
            "category": "የፀሐይ ጉልበት",
            "factline": "173,000 Terawatts • የውሃ ኡደት • ንጹህ የሶላር ኤሌክትሪክ",
            "img": "04-sun-power-tiktok.jpg",
            "hook": "“ፀሐይ ለምድር የምትሰጠው ኃይል የሰው ልጅ መጠቀም ከሚችለው በላይ ነው!”",
            "article": [
                "የናሳ (NASA EarthData) ሳይንሳዊ መረጃ እንደሚያሳየው ምድር ከፀሐይ የምታገኘው የኃይል መጠን የሰው ልጅ በሙሉ በዓመት ከሚጠቀመው ኤሌክትሪክና አጠቃላይ ኃይል ከ10,000 እጥፍ በላይ ይበልጣል!",
                "በእርግጥ ያንን ኃይል በሙሉ መሰብሰብ አንችልም ምክንያቱም ምድር ድቡልቡል ናት፣ አየርና ደመናዎች የተወሰነውን ያንጸባርቃሉ፣ እንዲሁም ግማሹ ምድር ሁልጊዜ በሌሊት ጨለማ ውስጥ ነው። ነገር ግን የሶላር ፓነል የፀሐይን ብርሃን በሙሉ መሰብሰብ አይጠበቅበትም፤ በእሱ ላይ የሚያርፈውን ትንሽ ክፍል ወደ ኤሌክትሪክ መለወጥ ብቻ በቂ ነው!"
            ],
            "script": [
                ("0–05 ሰከንድ", "ትንሽ ምድር ከፀሐይ አጠገብ።", "ፀሐይ ምን ያህል ኃይለኛ ናት?"),
                ("25–39 ሰከንድ", "10,000× የሚል ትልቅ ቁጥር።", "ናሳ ምድር ከፀሐይ የምታገኘው ኃይል ከምንጠቀመው ከ10,000 እጥፍ ይበልጣል ይላል።")
            ],
            "caption": "ምድር ከፀሐይ የምታገኘው ኃይል የሰው ልጅ ከሚጠቀመው ከ10,000 እጥፍ ይበልጣል። የሶላር ቴክኖሎጂ ተፈጥሮ የሰጠችንን ታላቅ ኃይል በጥበብ የመጠቀም መንገድ ነው!",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #ፀሐይ #ኢትዮጵያ #ሳይንስ"
        },
        {
            "num": "05",
            "title": "የሶላር ፓነሎች ለምን ሰማያዊ ወይም ጥቁር ይሆናሉ?",
            "category": "የዕለት ተዕለት ሳይንስ",
            "factline": "ፖሊክሪስታላይን ሰማያዊ • ሞኖክሪስታላይን ጥቁር • የነጸብራቅ ሽፋን",
            "img": "05-panel-colour-tiktok.jpg",
            "hook": "“የሶላር ፓነሎች ለምን ነጭ ወይም ቀይ ሳይሆኑ ሰማያዊና ጥቁር ይሆናሉ?”",
            "article": [
                "የሶላር ሴል የሚመጣውን ብርሃን በሙሉ መምጠጥ አለበት፤ ብርሃን ተመልሶ ከተንፀባረቀ ኤሌክትሪክ ማመንጨት አይችልም። ሲሊኮን በተፈጥሮው ብርሃን የሚያንፀባርቅ በመሆኑ፣ አምራቾች የብርሃን ነጸብራቅ መከላከያ (Anti-reflective coating) እና ልዩ ወለል ይጠቀማሉ—ይህም ፓነሎቹ ጥቁር ሰማያዊ ወይም ጥቁር እንዲመስሉ ያደርጋቸዋል።",
                "ብዙ ክሪስታሎች ያሏቸው (Multicrystalline) ሰማያዊ መልክ አላቸው፤ አንድ ወጥ ክሪስታል ያላቸው (Monocrystalline) ደግሞ ጥቁር መልክ አላቸው። ነገር ግን ቀለም ብቻውን የፓነል ጥራት መለኪያ አይደለም!"
            ],
            "script": [
                ("0–05 ሰከንድ", "ሰማያዊና ጥቁር ፓነሎች ጎን ለጎን።", "የትኛው ይበልጣል፡ ሰማያዊ ወይስ ጥቁር?"),
                ("40–53 ሰከንድ", "ቀለም = ጥራት የሚለው ላይ ቀይ X ምልክት።", "ነገር ግን ቀለም ብቻውን የፓነሉን ጥራት አይናገርም። በትክክለኛ ምህንድስና ይምረጡ!")
            ],
            "caption": "ሰማያዊና ጥቁር የሶላር ፓነሎች ብርሃን መምጠጥ እንዲችሉ የተሰሩ ናቸው። ቀለም ብቻውን የጥራት ማረጋገጫ ስላልሆነ በትክክለኛ ዝርዝርና ገጠማ ይምረጡ!",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #የሶላርፓነል #ኢትዮጵያ"
        },
        {
            "num": "06",
            "title": "ፀሐይ ድንገት ብትጠፋ ምን ይፈጠራል?",
            "category": "አሳቢ ጥያቄ",
            "factline": "ለ8 ደቂቃ ያህል • ጨለማ • ምድር ከመዞር ወጥታ በቀጥታ መጓዝ",
            "img": "06-sun-disappears-tiktok.jpg",
            "hook": "“ፀሐይ አሁን ድንገት ብትጠፋ ለ8 ደቂቃ ያህል ምንም አታውቁም!”",
            "article": [
                "ይህ በሳይንስ ሊከሰት የሚችል ነገር ባይሆንም ትልቅ የሳይንስ ጥያቄ ነው። የፀሐይ ብርሃን ወደ ምድር ለመድረስ 8 ደቂቃ ያህል ይወስዳል። የስበት ኃይል ለውጥም በብርሃን ፍጥነት ነው የሚጓዘው።",
                "ስለሆነም ፀሐይ ድንገት በቅጽበት ብትጠፋ፣ ምድር ለ8 ደቂቃ ያህል ብርሃን ማግኘቷንና በሕዋ ላይ መዞሯን ትቀጥላለች። ከ8 ደቂቃ በኋላ ግን ሰማዩ ይጨልማል፣ ምድርም ከመዞር ወጥታ በቀጥታ መስመር ትጓዛለች። የሶላር ፓነሎች ማመንጨት ያቆማሉ።"
            ],
            "script": [
                ("0–04 ሰከንድ", "መደበኛ ጠዋት። ፀሐይ ትጠፋለች።", "ፀሐይ አሁን ድንገት ብትጠፋ ለ8 ደቂቃ ያህል ምንም አታውቁም።"),
                ("04–14 ሰከንድ", "ሰዓት ወደ 8 ደቂቃ መቁጠር ይጀምራል።", "ወደ ምድር በመጓዝ ላይ ያለው ብርሃን ለ8 ደቂቃ መድረሱን ይቀጥላል።"),
                ("24–36 ሰከንድ", "በ8ኛው ደቂቃ አለም ጨለማ ትሆናለች።", "ከዚያ የመጨረሻው ብርሃን ያልፋል፣ ሰማዩም ይጨልማል።")
            ],
            "caption": "ፀሐይ በድንገት ብትጠፋ ምድር ለ8 ደቂቃ ያህል የነበረውን ብርሃን ታገኛለች። ከዚያ በኋላ ጨለማ ይሆናል።",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #ፀሐይ #ሳይንስ #ኢትዮጵያ"
        },
        {
            "num": "07",
            "title": "ከብሔራዊ የኤሌክትሪክ መስመር ውጭ መኖር ይችላሉ?",
            "category": "Off-Grid ፈተና",
            "factline": "የኃይል ፍላጎት ማወቅ • ባትሪ መደገፍ • የፀሐይ መጠን ማዘጋጀት",
            "img": "07-off-grid-tiktok.jpg",
            "hook": "“ዛሬ ከቆጣሪ ተላቀው በሶላር ብቻ መኖር ይችላሉ?”",
            "article": [
                "ከብሔራዊ መብራት ውጭ መኖር ቀላል ይመስላል፡ ፓነልና ባትሪ ገዝቶ ከመብራት መቆራረጥ እፎይ ማለት! ነገር ግን አስተማማኝ የሶላር ሲስተም የሚጀመረው በጥያቄዎች እንጂ እቃ በመግዛት አይደለም።",
                "ዋና ዋና እቃዎችን (መብራት፣ ዋይፋይ፣ ፍሪጅ) ከከበሩ እቃዎች (የኤሌክትሪክ ምድጃ፣ ቦይለር) ለይቶ ማወቅ፣ የባትሪ መጠንና የፓነል አቅምን በትክክል ማስላት ይጠይቃል።"
            ],
            "script": [
                ("0–05 ሰከንድ", "እጅ የመብራት ቆጣሪ ሲያጠፋ።", "ዛሬ ከቆጣሪ ተላቀው በሶላር ብቻ መኖር ይችላሉ?"),
                ("05–15 ሰከንድ", "ቤቱ በሶላር ሲበራ።", "ይቻላል—ነገር ግን ዝም ብሎ ባትሪ ከመግዛት ይበልጣል። ፍላጎትዎን በጥራት ያሰሉ!")
            ],
            "caption": "ከብሔራዊ መብራት ውጭ መኖር ይቻላል? ትክክለኛ የሶላር ሲስተም የሚጀመረው የቤት እቃዎችን የኃይል ፍላጎት በጥራት በማወቅ ነው።",
            "hashtags": "#ዩኒቨርሳልኤሌክትሮኒክስ #OffGridLiving #ኢትዮጵያ #አዲስአበባ"
        }
    ]

    for item in am_stories:
        h2 = doc.add_paragraph()
        run_h2 = h2.add_run(f"ታሪክ {item['num']}: {item['title']}")
        run_h2.font.name = 'Nyala'
        run_h2.font.size = Pt(17)
        run_h2.font.bold = True
        run_h2.font.color.rgb = RGBColor(11, 47, 143)

        p_meta = doc.add_paragraph()
        run_meta = p_meta.add_run(f"ምድብ: {item['category']} | {item['factline']}")
        run_meta.font.name = 'Nyala'
        run_meta.font.size = Pt(10.5)
        run_meta.font.bold = True
        run_meta.font.color.rgb = RGBColor(22, 73, 198)

        # Image Insertion
        img_path = os.path.join(assets_dir, item['img'])
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(12)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(2.5))

        # Hook Box Table
        table_hook = doc.add_table(rows=1, cols=1)
        table_hook.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_hook = table_hook.cell(0, 0)
        set_cell_background(cell_hook, "FFF9ED")
        p_hook = cell_hook.paragraphs[0]
        p_hook.paragraph_format.space_before = Pt(6)
        p_hook.paragraph_format.space_after = Pt(6)
        run_hk_lbl = p_hook.add_run("የቲክቶክ መግቢያ (OPENING HOOK):\n")
        run_hk_lbl.font.name = 'Nyala'
        run_hk_lbl.font.size = Pt(9.5)
        run_hk_lbl.font.bold = True
        run_hk_lbl.font.color.rgb = RGBColor(179, 123, 0)
        run_hk_txt = p_hook.add_run(item['hook'])
        run_hk_txt.font.name = 'Nyala'
        run_hk_txt.font.size = Pt(12)
        run_hk_txt.font.bold = True
        run_hk_txt.font.color.rgb = RGBColor(6, 24, 69)

        doc.add_paragraph() # space

        # Article Paragraphs
        for p_txt in item['article']:
            p_art = doc.add_paragraph()
            r_art = p_art.add_run(p_txt)
            r_art.font.name = 'Nyala'
            r_art.font.size = Pt(11.5)
            p_art.paragraph_format.space_after = Pt(6)

        # Script Table
        p_sc_head = doc.add_paragraph()
        r_sc_head = p_sc_head.add_run("📱 90-ሰከንድ የአማርኛ ቲክቶክ ስክሪፕት (Teleprompter)")
        r_sc_head.font.name = 'Nyala'
        r_sc_head.font.bold = True
        r_sc_head.font.size = Pt(12)
        r_sc_head.font.color.rgb = RGBColor(11, 47, 143)

        table_sc = doc.add_table(rows=1, cols=3)
        table_sc.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table_sc.rows[0].cells
        hdr_cells[0].text = "ጊዜ (Timing)"
        hdr_cells[1].text = "የቪዲዮ እይታ (Visual)"
        hdr_cells[2].text = "የድምፅ ንግግር (Narration)"

        for cell in hdr_cells:
            set_cell_background(cell, "EEF3FF")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Nyala'
                    r.font.bold = True
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = RGBColor(6, 24, 69)

        for step in item['script']:
            row_cells = table_sc.add_row().cells
            row_cells[0].text = step[0]
            row_cells[1].text = step[1]
            row_cells[2].text = f'"{step[2]}"'
            for idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Nyala'
                        r.font.size = Pt(10)
                        if idx == 0:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(11, 47, 143)
                        elif idx == 2:
                            r.font.bold = True

        doc.add_paragraph()

        # Caption
        p_cap = doc.add_paragraph()
        r_cap_lbl = p_cap.add_run("የሶሻል ሚዲያ መግለጫና ሃሽታጎች:\n")
        r_cap_lbl.font.name = 'Nyala'
        r_cap_lbl.font.bold = True
        r_cap_lbl.font.size = Pt(10.5)
        r_cap_txt = p_cap.add_run(item['caption'] + "\n" + item['hashtags'])
        r_cap_txt.font.name = 'Nyala'
        r_cap_txt.font.size = Pt(10.5)
        r_cap_txt.font.italic = True

        doc.add_paragraph("--------------------------------------------------------------------------------")

    output_path = os.path.join(base_dir, "Universal_Solar_Stories_Amharic.docx")
    doc.save(output_path)
    print("Created:", output_path)

if __name__ == "__main__":
    create_english_doc()
    create_amharic_doc()
